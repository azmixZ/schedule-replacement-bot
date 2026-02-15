import pandas as pd
from datetime import datetime
from docx import Document
import os
import re

# пути и настройки
CURRENT_DIR = os.path.dirname(os.path.abspath(__file__))
BASE_DIR = os.path.dirname(CURRENT_DIR)
# пути к папке data
WORD_FILE = os.path.join(BASE_DIR, "data", "zameni.docx")
EXCEL_FILE = os.path.join(BASE_DIR, "data", "raspisanie.xlsx")

days_map = {
    "Monday": "Понедельник",
    "Tuesday": "Вторник",
    "Wednesday": "Среда",
    "Thursday": "Четверг",
    "Friday": "Пятница",
    "Saturday": "Суббота",
    "Sunday": "Воскресенье"
}

# фукнции обработки

def normalize_group(name: str) -> str:
    """Убирает пробелы и приводит к верхнему регистру."""
    if not isinstance(name, str):
        return str(name)
    return name.strip().upper().replace("\n", "").replace(" ", "")

def clean_pair_number(text: str) -> str:
    """Извлекает только цифру из номера пары."""
    digits = re.findall(r'\d+', str(text))
    return digits[0] if digits else ""

def check_replacements_status():
    """Проверяет наличие замен в файле Word."""
    if not os.path.exists(WORD_FILE):
        return "❌ Файл замен не найден."
    try:
        doc = Document(WORD_FILE)
        if len(doc.tables) > 0 and len(doc.tables[0].rows) > 1:
            return " Замены найдены и применены."
        return "️ Замен на сегодня нет."
    except Exception:
        return " Ошибка чтения файла замен."

# проверка замен

def read_replacements(target_day_name):
    replacements = {}
    if not os.path.exists(WORD_FILE):
        return replacements

    try:
        doc = Document(WORD_FILE)
        if not doc.paragraphs or not doc.tables:
            return replacements

        # дата с первой строк ворда
        date_text = doc.paragraphs[0].text.strip().lower()
        print(f"[DEBUG] Строка с датой в Word: '{date_text}'") # отладка прост

        # месяцы
        months = {
            "января": 1, "январь": 1,
            "февраля": 2, "февраль": 2,
            "марта": 3, "март": 3,
            "апреля": 4, "апрель": 4,
            "мая": 5, "май": 5,
            "июня": 6, "июнь": 6,
            "июля": 7, "июль": 7,
            "августа": 8, "август": 8,
            "сентября": 9, "сентябрь": 9,
            "октября": 10, "октябрь": 10,
            "ноября": 11, "ноябрь": 11,
            "декабря": 12, "декабрь": 12
        }

        # лишние символы убираем
        clean_text = date_text.replace("г.", "").replace(".", "").replace(",", "").strip()
        parts = clean_text.split()

        if len(parts) >= 3:
            try:
                day = int(parts[0])
                month_str = parts[1]
                month = months.get(month_str)
                year = int(parts[2])

                if month is None:
                    print(f"[WARN] Месяц '{month_str}' не найден в словаре.")
                    return {}

                # проверка ДН
                dt = datetime(year, month, day)
                word_day_name = days_map.get(dt.strftime("%A"))
                
                print(f"[DEBUG] Дата определена: {dt.date()} ({word_day_name})")

                # отладка прост на случай отсутвися замен
                if word_day_name != target_day_name:
                    print(f"[INFO] Замены в Word для {word_day_name}, а нам нужен {target_day_name}. Пропускаем.")
                    return {}
            except ValueError:
                print(f"[ERROR] Не удалось преобразовать числа в дате: {parts}")
                return {}
        else:
            print(f"[WARN] Недостаточно данных в строке даты: {parts}")
            return {}

    except Exception as e:
        print(f"[ERROR] Ошибка при проверке даты в Word: {e}")
        return {}

    # проверка таблцы через дату
    table = doc.tables[0]
    for row in table.rows[1:]:
        cells = row.cells
        if len(cells) < 7: continue

        pair_num = clean_pair_number(cells[1].text)
        group = normalize_group(cells[2].text)
        
        # сборка пары
        subject = cells[5].text.strip()
        room = cells[6].text.strip()
        
        if pair_num and group and subject:
            replacements[(group, pair_num)] = f"{subject} — {room}"
            
    return replacements
import re

def compact_lesson_name(text):
    """Оставляет только Преподавателя и место проведения (кабинет или зал)"""
    if not text:
        return "---"
    
    # лишние пометки убираем
    text = text.replace("🔄 (Замена)", "").strip()
    
    # поиск препода через символы
    teacher_match = re.search(r'([А-Я][а-я]+\s+[А-Я]\.[А-Я]\.)', text)
    teacher = teacher_match.group(1) if teacher_match else ""
    
    # проверка
    # физра
    gym_match = re.search(r'(спортзал|дискозал|тренаж\w*\s*зал|тренажерка|с/зал)', text, re.IGNORECASE)
    
    # кабинет
    room_match = re.search(r'(\d+)\s*(?:каб|к|кабинет)', text, re.IGNORECASE)
    
    location = ""
    if gym_match:
        location = gym_match.group(1).lower()
    elif room_match:
        location = f"{room_match.group(1)} каб."

    # результат
    if teacher and location:
        return f"{teacher} ({location})"
    elif teacher:
        return teacher
    elif location:
        return location.capitalize()
    
    # если не нащло
    return text if len(text) < 30 else text[:27] + "..."

def get_schedule(user_group: str, target_day: str = None):
    MY_GROUP = normalize_group(user_group)
    
    if not os.path.exists(EXCEL_FILE):
        return "Ошибка: Файл Excel не найден."

    try:
        df = pd.read_excel(EXCEL_FILE, header=None)
        df.iloc[:, 0] = df.iloc[:, 0].ffill()
        df.iloc[:, 1] = df.iloc[:, 1].ffill()
    except Exception as e:
        return f"Ошибка чтения Excel: {e}"

    # группа
    group_column = None
    header_row = df.iloc[0]
    for col_idx, cell_value in enumerate(header_row):
        if normalize_group(str(cell_value)) == MY_GROUP:
            group_column = col_idx
            break

    if group_column is None:
        return f"Группа {user_group} не найдена."

    # день
    if target_day and target_day in days_map.values():
        current_day = target_day
    else:
        current_day = days_map.get(datetime.now().strftime("%A"), "Понедельник")
    
    replacements = read_replacements(current_day)

    # расписание 
    start_idx = None
    for idx, row in df.iterrows():
        if current_day.lower() in str(row[0]).lower():
            start_idx = idx
            break
            
    if start_idx is None:
        return f"Расписание на {current_day} в Excel не найдено."

    schedule_dict = {}
    for i in range(0, 15):
        c_idx = start_idx + i
        if c_idx >= len(df): break
        row = df.iloc[c_idx]
        
        if i > 0:
            day_cell = str(row[0]).strip().lower()
            if day_cell != current_day.lower() and any(d.lower() in day_cell for d in days_map.values()):
                break

        p_num = clean_pair_number(str(row[1]))
        if not p_num: continue

        lesson_excel = str(row[group_column]).strip() if pd.notna(row[group_column]) else ""
        if lesson_excel.lower() in ["nan", ""]: lesson_excel = ""

        key = (MY_GROUP, p_num)
        if key in replacements:
            # спецсимвол вцввода
            schedule_dict[p_num] = (replacements[key], True)
        elif lesson_excel:
            schedule_dict[p_num] = (lesson_excel, False)

    # если пары нет то замена всё равно наложится
    for (g, p), lesson in replacements.items():
        if g == MY_GROUP and p not in schedule_dict:
            schedule_dict[p] = (lesson, True)

    if not schedule_dict:
        return f"📅 {current_day}\nДля группы {MY_GROUP} пар не найдено."

    # краткий вывод
    msg = [f"📅 <b>{current_day}</b> | {MY_GROUP}", "────────────────────"]
    
    for p_num in sorted(schedule_dict.keys(), key=lambda x: int(x)):
        lesson_text, is_repl = schedule_dict[p_num]
        
        # сокращение
        short_name = compact_lesson_name(lesson_text)
        
        icon = " 🔄" if is_repl else ""
        msg.append(f"<b>{p_num} п:</b> {short_name}{icon}")

    return "\n".join(msg)